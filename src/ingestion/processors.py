import os
import re
from pathlib import Path
from typing import List, Optional, Union
import pypdf
import docx
from haystack import Document
import io
import logging

from src.config import LateChunkingConfig

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text(self, file_source: Path | str | bytes, filename: Optional[str] = None) -> str:
        """Extracts text from PDF or DOCX. Source can be a path or bytes."""
        import io
        
        if isinstance(file_source, (Path, str)):
            file_path = Path(file_source)
            suffix = file_path.suffix.lower()
            with open(file_path, "rb") as f:
                stream = io.BytesIO(f.read())
        elif isinstance(file_source, bytes):
            stream = io.BytesIO(file_source)
            if not filename:
                # We need a filename or at least an extension to know how to process
                raise ValueError("filename must be provided when passing bytes to extract_text")
            suffix = Path(filename).suffix.lower()
        else:
            raise ValueError("Unsupported file source type")

        if suffix == ".pdf":
            return self._extract_pdf(stream)
        elif suffix == ".docx":
            return self._extract_docx(stream)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _extract_pdf(self, stream: io.BytesIO) -> str:
        text = ""
        reader = pypdf.PdfReader(stream)
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_docx(self, stream: io.BytesIO) -> str:
        doc = docx.Document(stream)
        return "\n".join([para.text for para in doc.paragraphs])

    def create_chunks(self, text: str, metadata: dict) -> List[Document]:
        """Simple chunking logic with overlap."""
        # Clean text
        text = re.sub(r'\s+', ' ', text).strip()
        
        words = text.split()
        chunks = []
        
        # Use a simple word-based chunking for efficiency
        # In a real scenario, you might want to use Haystack's DocumentSplitter
        # but here we implement a manual one as requested for flexibility
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            if chunk_text:
                chunks.append(Document(content=chunk_text, meta=metadata.copy()))
        
        return chunks

    @staticmethod
    def infer_campaign(filename: str) -> str:
        """Infers campaign name from filename (e.g., 'Summer2024_promo.pdf' -> 'Summer2024')."""
        match = re.search(r'^([^_.-]+)', filename)
        return match.group(1) if match else "General"


class LateChunkerProcessor:
    """Document processor that uses Late Chunking for context-preserving embeddings.

    This class wraps :class:`~src.indexing.late_chunker.LateChunker` and exposes
    the same ``create_chunks`` / ``extract_text`` interface as
    :class:`DocumentProcessor` so it can be used as a drop-in replacement.

    Unlike ``DocumentProcessor``, the ``Document`` objects returned by
    ``create_chunks`` already have their ``embedding`` field populated.

    Args:
        config: Optional :class:`LateChunkingConfig`. Defaults are used when
            ``None``.
    """

    def __init__(self, config: Optional[LateChunkingConfig] = None) -> None:
        self.config = config or LateChunkingConfig()
        # Lazy import to avoid heavy model loading at module import time
        from src.indexing.late_chunker import LateChunker

        self._chunker = LateChunker(self.config)
        # Reuse DocumentProcessor for text extraction only
        self._text_extractor = DocumentProcessor()

    def extract_text(self, file_source: Path | str | bytes, filename: Optional[str] = None) -> str:
        """Extract text from a file. Delegates to :class:`DocumentProcessor`."""
        return self._text_extractor.extract_text(file_source, filename=filename)

    def create_chunks(self, text: str, metadata: dict) -> List[Document]:
        """Split *text* using Late Chunking and return documents with embeddings.

        Each returned :class:`~haystack.Document` has its ``embedding``
        field already set, so no separate embedding step is needed in the
        indexing pipeline.

        Args:
            text: The full document text.
            metadata: Base metadata to attach to every chunk.

        Returns:
            List of Haystack ``Document`` objects with pre-computed embeddings.
        """
        return self._chunker.process_document(text, metadata)

    @staticmethod
    def infer_campaign(filename: str) -> str:
        """Infers campaign name from filename. Same logic as DocumentProcessor."""
        return DocumentProcessor.infer_campaign(filename)


def get_processor(
    use_late_chunking: bool = False,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    late_chunking_config: Optional[LateChunkingConfig] = None,
) -> Union[DocumentProcessor, LateChunkerProcessor]:
    """Factory function that returns the appropriate document processor.

    Args:
        use_late_chunking: When ``True``, return a
            :class:`LateChunkerProcessor`; otherwise return a
            :class:`DocumentProcessor`.
        chunk_size: Word-based chunk size (only for ``DocumentProcessor``).
        chunk_overlap: Word overlap (only for ``DocumentProcessor``).
        late_chunking_config: Config for late chunking (only for
            ``LateChunkerProcessor``).

    Returns:
        A processor instance with ``extract_text`` and ``create_chunks``
        methods.
    """
    if use_late_chunking:
        logger.info("Using LateChunkerProcessor")
        return LateChunkerProcessor(config=late_chunking_config)
    else:
        logger.info("Using standard DocumentProcessor")
        return DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
